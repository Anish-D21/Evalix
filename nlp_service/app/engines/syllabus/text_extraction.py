"""
Text extraction engine (spec Section 14).

Extracts raw text from an uploaded syllabus file entirely in memory --
no temporary files are written to disk, matching Section 48's
requirement that the core system not depend on persistent filesystem
storage. Supports the three formats the spec calls out: PDF (PyMuPDF),
DOCX (python-docx), and plain TXT.

Extraction is intentionally "dumb" here: it does not try to guess
structure. Unit/topic detection is a separate, later stage
(engines/syllabus/unit_detection.py) -- this module's only job is
"bytes in, plain text out."
"""

from __future__ import annotations

import io


class UnsupportedFileTypeError(ValueError):
    """Raised when the uploaded file isn't PDF, DOCX, or TXT."""


class TextExtractionError(ValueError):
    """Raised when a file of a supported type fails to parse (e.g. corrupt PDF)."""


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def _extract_pdf_text(file_bytes: bytes) -> str:
    import fitz  # PyMuPDF

    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            pages = [page.get_text() for page in doc]
    except Exception as exc:  # noqa: BLE001 - surface as a clean extraction error
        raise TextExtractionError(f"Could not parse PDF: {exc}") from exc
    return "\n".join(pages)


def _extract_docx_text(file_bytes: bytes) -> str:
    import docx

    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as exc:  # noqa: BLE001
        raise TextExtractionError(f"Could not parse DOCX: {exc}") from exc

    parts = [p.text for p in document.paragraphs]
    # Tables are common in syllabus documents (e.g. a unit/topic/hours
    # table) -- pull their cell text too rather than silently dropping it.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_txt_text(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise TextExtractionError("Could not decode text file with any supported encoding.")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Dispatches to the right extractor based on the file extension.
    Raises UnsupportedFileTypeError for anything other than
    .pdf / .docx / .txt, and TextExtractionError if the file matches a
    supported extension but fails to parse (e.g. corrupted upload).
    """
    lower = filename.lower()

    if lower.endswith(".pdf"):
        return _extract_pdf_text(file_bytes)
    if lower.endswith(".docx"):
        return _extract_docx_text(file_bytes)
    if lower.endswith(".txt"):
        return _extract_txt_text(file_bytes)

    raise UnsupportedFileTypeError(
        f"Unsupported file type for '{filename}'. Evalix accepts PDF, DOCX, and TXT syllabus files."
    )
