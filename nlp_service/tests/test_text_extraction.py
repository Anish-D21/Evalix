import pytest

from app.engines.syllabus.text_extraction import TextExtractionError, UnsupportedFileTypeError, extract_text


def _build_pdf_bytes(text_lines):
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    y = 72
    for line in text_lines:
        page.insert_text((72, y), line)
        y += 20
    data = doc.tobytes()
    doc.close()
    return data


def _build_docx_bytes(paragraphs):
    import io

    import docx

    d = docx.Document()
    for p in paragraphs:
        d.add_paragraph(p)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_extract_txt():
    text = extract_text(b"Unit 1: Test\nSome content here.", "syllabus.txt")
    assert "Unit 1: Test" in text
    assert "Some content here." in text


def test_extract_txt_latin1_fallback():
    # bytes that aren't valid UTF-8 but are valid latin-1
    raw = "café résumé".encode("latin-1")
    text = extract_text(raw, "notes.txt")
    assert "caf" in text


def test_extract_pdf():
    pdf_bytes = _build_pdf_bytes(["Unit 1: Data Structures", "Arrays and linked lists."])
    text = extract_text(pdf_bytes, "syllabus.pdf")
    assert "Data Structures" in text
    assert "Arrays" in text


def test_extract_pdf_case_insensitive_extension():
    pdf_bytes = _build_pdf_bytes(["Unit 1: Data Structures"])
    text = extract_text(pdf_bytes, "syllabus.PDF")
    assert "Data Structures" in text


def test_extract_docx():
    docx_bytes = _build_docx_bytes(["Unit 1: Data Structures", "Arrays and linked lists."])
    text = extract_text(docx_bytes, "syllabus.docx")
    assert "Data Structures" in text
    assert "Arrays" in text


def test_extract_docx_includes_table_content():
    import io

    import docx

    d = docx.Document()
    d.add_paragraph("Unit 1: Overview")
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Topic"
    table.rows[0].cells[1].text = "Recursion"
    buf = io.BytesIO()
    d.save(buf)

    text = extract_text(buf.getvalue(), "syllabus.docx")
    assert "Recursion" in text


def test_unsupported_file_type_rejected():
    with pytest.raises(UnsupportedFileTypeError):
        extract_text(b"whatever", "syllabus.exe")


def test_corrupt_pdf_raises_extraction_error():
    with pytest.raises(TextExtractionError):
        extract_text(b"this is not a real pdf", "syllabus.pdf")


def test_corrupt_docx_raises_extraction_error():
    with pytest.raises(TextExtractionError):
        extract_text(b"this is not a real docx", "syllabus.docx")
